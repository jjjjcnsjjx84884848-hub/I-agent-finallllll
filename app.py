from __future__ import annotations

import csv
import io
import json
import re
from datetime import datetime, timezone
from pathlib import Path
from typing import Any

import streamlit as st

import ai
from scanner import detect_language, extract_budget, score_job, stable_id


ROOT = Path(__file__).resolve().parent
CONFIG_PATH = ROOT / "config.json"
JOBS_PATH = ROOT / "jobs.json"
STATUS_PATH = ROOT / "scan_status.json"
FAVORITES_PATH = ROOT / "favorites.json"
HISTORY_PATH = ROOT / "history.json"

st.set_page_config(
    page_title="Фриланс-Агент",
    page_icon="✨",
    layout="wide",
    initial_sidebar_state="expanded",
)


# ---------------------------------------------------------
# СТИЛИ ИНТЕРФЕЙСА
# ---------------------------------------------------------

st.markdown(
    """
    <style>
        .block-container {
            max-width: 1450px;
            padding-top: 1.4rem;
            padding-bottom: 4rem;
        }

        .main-title {
            font-size: 2.2rem;
            font-weight: 800;
            margin-bottom: 0.2rem;
        }

        .main-subtitle {
            opacity: 0.78;
            margin-bottom: 1.4rem;
        }

        .hero-card {
            border: 1px solid rgba(128, 128, 128, 0.25);
            border-radius: 22px;
            padding: 1.4rem 1.5rem;
            margin-bottom: 1.2rem;
            background: rgba(127, 127, 127, 0.06);
        }

        .job-card {
            border: 1px solid rgba(128, 128, 128, 0.24);
            border-radius: 18px;
            padding: 1rem 1.1rem;
            margin-bottom: 0.85rem;
            background: rgba(127, 127, 127, 0.045);
        }

        .job-title {
            font-size: 1.05rem;
            font-weight: 750;
            margin-bottom: 0.35rem;
        }

        .job-meta {
            opacity: 0.72;
            font-size: 0.88rem;
            margin-bottom: 0.6rem;
        }

        .badge {
            display: inline-block;
            border-radius: 999px;
            padding: 0.22rem 0.58rem;
            margin-right: 0.35rem;
            margin-bottom: 0.35rem;
            font-size: 0.78rem;
            border: 1px solid rgba(128, 128, 128, 0.28);
        }

        .risk-low {
            background: rgba(50, 180, 90, 0.13);
        }

        .risk-medium {
            background: rgba(230, 170, 40, 0.14);
        }

        .risk-high {
            background: rgba(220, 65, 65, 0.14);
        }

        .simple-task {
            background: rgba(80, 120, 255, 0.13);
        }

        .copy-box {
            border: 1px dashed rgba(128, 128, 128, 0.42);
            border-radius: 14px;
            padding: 1rem;
            white-space: pre-wrap;
            background: rgba(127, 127, 127, 0.04);
        }

        div[data-testid="stMetric"] {
            border: 1px solid rgba(128, 128, 128, 0.20);
            border-radius: 17px;
            padding: 0.8rem 1rem;
            background: rgba(127, 127, 127, 0.04);
        }

        div[data-testid="stExpander"] {
            border-radius: 16px;
            overflow: hidden;
        }

        .small-muted {
            opacity: 0.67;
            font-size: 0.85rem;
        }
    </style>
    """,
    unsafe_allow_html=True,
)


# ---------------------------------------------------------
# КОНСТАНТЫ
# ---------------------------------------------------------

RISK_LABELS = {
    "low": "Низкий",
    "medium": "Средний",
    "high": "Высокий",
}

RISK_ICONS = {
    "low": "🟢",
    "medium": "🟡",
    "high": "🔴",
}

LANGUAGE_LABELS = {
    "ru": "Русский",
    "uk": "Украинский",
    "de": "Немецкий",
    "en": "Английский",
    "auto": "Определить автоматически",
}

TASK_LABELS = {
    "translation": "Перевод",
    "transcription": "Транскрибация",
    "data_entry": "Ввод данных",
    "proofreading": "Проверка текста",
    "virtual_assistant": "Виртуальный помощник",
    "writing": "Работа с текстом",
    "other": "Другое",
}

SIMPLE_TASK_KEYWORDS = [
    "translation",
    "translate",
    "translator",
    "proofreading",
    "proofread",
    "proofreader",
    "transcription",
    "transcribe",
    "data entry",
    "copy paste",
    "copy-paste",
    "typing",
    "retyping",
    "subtitle",
    "subtitles",
    "captioning",
    "text editing",
    "document formatting",
    "pdf to word",
    "image to text",
    "audio to text",
    "web research",
    "internet research",
    "virtual assistant",
    "product description",
    "short description",
    "email writing",
    "localization",
    "übersetzung",
    "übersetzen",
    "übersetzer",
    "korrekturlesen",
    "korrektur",
    "transkription",
    "dateneingabe",
    "untertitel",
    "texterfassung",
    "internetrecherche",
    "virtuelle assistenz",
    "перевод",
    "перевести",
    "переводчик",
    "проверка текста",
    "редактирование текста",
    "расшифровка",
    "транскрибация",
    "набор текста",
    "ввод данных",
    "копирование данных",
    "субтитры",
    "описание товара",
    "поиск информации",
    "виртуальный помощник",
    "переклад",
    "перекласти",
    "перекладач",
    "перевірка тексту",
    "редагування тексту",
    "розшифровка",
    "транскрипція",
    "набір тексту",
    "введення даних",
    "копіювання даних",
    "субтитри",
    "опис товару",
    "пошук інформації",
    "віртуальний помічник",
]

COMPLEX_TASK_KEYWORDS = [
    "senior developer",
    "senior engineer",
    "lead developer",
    "full stack",
    "full-stack",
    "backend developer",
    "frontend developer",
    "mobile application",
    "machine learning",
    "blockchain",
    "devops",
    "cybersecurity",
    "penetration testing",
    "software architecture",
    "3d animation",
    "3d modeling",
    "legal advice",
    "medical diagnosis",
    "licensed professional",
    "5+ years",
    "five years experience",
]

SUPPORTED_TEXT_EXTENSIONS = {
    ".txt",
    ".md",
    ".csv",
    ".json",
}

SUPPORTED_DOCUMENT_EXTENSIONS = {
    ".docx",
    ".pdf",
}

MAX_FILE_TEXT_LENGTH = 100_000


# ---------------------------------------------------------
# БАЗОВЫЕ ФУНКЦИИ ДЛЯ ФАЙЛОВ JSON
# ---------------------------------------------------------

def utc_now() -> str:
    return (
        datetime.now(timezone.utc)
        .replace(microsecond=0)
        .isoformat()
    )


def load_json(path: Path, default: Any) -> Any:
    try:
        return json.loads(path.read_text(encoding="utf-8"))
    except Exception:
        return default


def save_json(path: Path, data: Any) -> bool:
    try:
        temporary_path = path.with_suffix(path.suffix + ".tmp")

        temporary_path.write_text(
            json.dumps(
                data,
                ensure_ascii=False,
                indent=2,
            ),
            encoding="utf-8",
        )

        temporary_path.replace(path)
        return True

    except Exception:
        return False


def get_config() -> dict[str, Any]:
    data = load_json(CONFIG_PATH, {})

    if isinstance(data, dict):
        return data

    return {}


def get_jobs() -> list[dict[str, Any]]:
    data = load_json(JOBS_PATH, [])

    if isinstance(data, list):
        return [
            job
            for job in data
            if isinstance(job, dict)
        ]

    return []


def get_favorites() -> list[str]:
    data = load_json(FAVORITES_PATH, [])

    if isinstance(data, list):
        return [
            str(item)
            for item in data
        ]

    return []


def get_history() -> list[dict[str, Any]]:
    data = load_json(HISTORY_PATH, [])

    if isinstance(data, list):
        return [
            item
            for item in data
            if isinstance(item, dict)
        ]

    return []


def add_history_entry(
    action_type: str,
    title: str,
    content: str,
    job_id: str = "",
) -> bool:
    history = get_history()

    history.insert(
        0,
        {
            "id": stable_id(
                action_type,
                utc_now(),
                title,
            ),
            "type": action_type,
            "title": title,
            "content": content,
            "job_id": job_id,
            "created_at": utc_now(),
        },
    )

    return save_json(
        HISTORY_PATH,
        history[:300],
    )


# ---------------------------------------------------------
# ФУНКЦИИ ДЛЯ ОЦЕНКИ И ФИЛЬТРАЦИИ ЗАДАНИЙ
# ---------------------------------------------------------

def job_full_text(job: dict[str, Any]) -> str:
    return (
        f"{job.get('title', '')} "
        f"{job.get('description', '')} "
        f"{job.get('category', '')} "
        f"{job.get('matched_keywords', '')}"
    ).lower()


def simple_task_score(job: dict[str, Any]) -> int:
    text = job_full_text(job)
    score = 0

    for keyword in SIMPLE_TASK_KEYWORDS:
        if keyword in text:
            score += 2

    for keyword in COMPLEX_TASK_KEYWORDS:
        if keyword in text:
            score -= 4

    description_length = len(
        str(job.get("description", ""))
    )

    if description_length < 2500:
        score += 1

    if description_length > 6500:
        score -= 2

    task_type = str(
        job.get("task_type", "")
    ).lower()

    if task_type in {
        "translation",
        "transcription",
        "data_entry",
        "proofreading",
    }:
        score += 4

    return score


def is_simple_task(job: dict[str, Any]) -> bool:
    stored_value = job.get("is_simple")

    if isinstance(stored_value, bool):
        return stored_value or simple_task_score(job) >= 1

    return simple_task_score(job) >= 1


def classify_task_local(job: dict[str, Any]) -> str:
    stored_type = str(
        job.get("task_type", "")
    ).strip()

    if stored_type and stored_type != "other":
        return stored_type

    text = job_full_text(job)

    categories = {
        "translation": [
            "translation",
            "translate",
            "translator",
            "localization",
            "übersetzung",
            "übersetzen",
            "перевод",
            "перевести",
            "переклад",
            "перекласти",
        ],
        "transcription": [
            "transcription",
            "transcribe",
            "audio to text",
            "расшифровка",
            "транскрибация",
            "transkription",
            "розшифровка",
            "транскрипція",
        ],
        "data_entry": [
            "data entry",
            "copy paste",
            "typing",
            "retyping",
            "dateneingabe",
            "ввод данных",
            "набор текста",
            "введення даних",
            "набір тексту",
        ],
        "proofreading": [
            "proofreading",
            "proofread",
            "text correction",
            "korrekturlesen",
            "korrektur",
            "проверка текста",
            "редактирование текста",
            "перевірка тексту",
            "редагування тексту",
        ],
        "virtual_assistant": [
            "virtual assistant",
            "virtuelle assistenz",
            "виртуальный помощник",
            "віртуальний помічник",
        ],
        "writing": [
            "writing",
            "copywriting",
            "product description",
            "email writing",
            "описание товара",
            "опис товару",
        ],
    }

    for task_type, keywords in categories.items():
        if any(
            keyword in text
            for keyword in keywords
        ):
            return task_type

    return "other"


def parse_budget_values(job: dict[str, Any]) -> list[float]:
    budget_text = str(
        job.get("budget_text", "")
    ).strip()

    if not budget_text:
        return []

    if budget_text.lower() in {
        "not stated",
        "unknown",
        "none",
        "не указан",
        "—",
    }:
        return []

    normalized = (
        budget_text
        .replace(",", ".")
        .replace(" ", "")
    )

    numbers = re.findall(
        r"\d+(?:\.\d+)?",
        normalized,
    )

    values: list[float] = []

    for number in numbers:
        try:
            value = float(number)

            if 0 < value < 100_000:
                values.append(value)

        except ValueError:
            continue

    return values


def budget_in_range(
    job: dict[str, Any],
    minimum: float,
    maximum: float,
    include_without_budget: bool,
) -> bool:
    values = parse_budget_values(job)

    if not values:
        return include_without_budget

    job_minimum = min(values)
    job_maximum = max(values)

    return (
        job_maximum >= minimum
        and job_minimum <= maximum
    )


def budget_display(job: dict[str, Any]) -> str:
    budget = str(
        job.get("budget_text", "")
    ).strip()

    if not budget or budget.lower() in {
        "not stated",
        "unknown",
        "none",
    }:
        return "Не указан"

    return budget


def risk_display(job: dict[str, Any]) -> str:
    risk = str(
        job.get("risk", "medium")
    ).lower()

    return RISK_LABELS.get(
        risk,
        "Не определён",
    )


def risk_icon(job: dict[str, Any]) -> str:
    risk = str(
        job.get("risk", "medium")
    ).lower()

    return RISK_ICONS.get(
        risk,
        "⚪",
    )


def language_display(job: dict[str, Any]) -> str:
    language = str(
        job.get("language_hint", "")
    ).lower()

    return LANGUAGE_LABELS.get(
        language,
        language or "Не определён",
    )


def task_display(job: dict[str, Any]) -> str:
    task_type = classify_task_local(job)

    return TASK_LABELS.get(
        task_type,
        "Другое",
    )


def calculate_opportunity_score(
    job: dict[str, Any],
) -> int:
    base_score = int(
        job.get("score", 0)
    )

    result = base_score

    if is_simple_task(job):
        result += 8

    if str(job.get("risk")) == "low":
        result += 7

    if str(job.get("risk")) == "high":
        result -= 25

    budget_values = parse_budget_values(job)

    if budget_values:
        budget_min = min(budget_values)
        budget_max = max(budget_values)

        if budget_max >= 5 and budget_min <= 50:
            result += 8

    if classify_task_local(job) in {
        "translation",
        "proofreading",
        "transcription",
        "data_entry",
    }:
        result += 7

    return max(
        0,
        min(100, result),
    )


def estimated_difficulty(job: dict[str, Any]) -> str:
    score = simple_task_score(job)
    description_length = len(
        str(job.get("description", ""))
    )

    if score >= 6 and description_length < 3000:
        return "Лёгкая"

    if score >= 1:
        return "Средняя"

    return "Сложная"


def estimated_time(job: dict[str, Any]) -> str:
    task_type = classify_task_local(job)
    description_length = len(
        str(job.get("description", ""))
    )

    base_minutes = {
        "translation": 45,
        "proofreading": 35,
        "transcription": 60,
        "data_entry": 40,
        "virtual_assistant": 75,
        "writing": 60,
        "other": 90,
    }.get(task_type, 90)

    if description_length > 4000:
        base_minutes += 45

    if description_length > 8000:
        base_minutes += 60

    if base_minutes < 60:
        return f"Около {base_minutes} мин."

    hours = round(base_minutes / 60, 1)
    return f"Около {hours} ч."


def recommended_price(job: dict[str, Any]) -> str:
    budget_values = parse_budget_values(job)

    if budget_values:
        minimum = min(budget_values)
        maximum = max(budget_values)

        if len(budget_values) >= 2:
            suggested = round(
                minimum + (maximum - minimum) * 0.45,
                2,
            )
        else:
            suggested = maximum

        return f"{suggested:g}"

    task_type = classify_task_local(job)

    suggested_prices = {
        "translation": 15,
        "proofreading": 12,
        "transcription": 18,
        "data_entry": 12,
        "virtual_assistant": 20,
        "writing": 18,
        "other": 20,
    }

    return str(
        suggested_prices.get(task_type, 20)
    )


def job_option_label(job: dict[str, Any]) -> str:
    score = calculate_opportunity_score(job)
    title = str(
        job.get("title", "Без названия")
    )[:85]

    return (
        f"{score}/100 · "
        f"{risk_icon(job)} {risk_display(job)} · "
        f"{task_display(job)} · "
        f"{title}"
    )


# ---------------------------------------------------------
# ИЗБРАННОЕ
# ---------------------------------------------------------

def is_favorite(job_id: str) -> bool:
    return job_id in set(get_favorites())


def toggle_favorite(job_id: str) -> bool:
    favorites = get_favorites()

    if job_id in favorites:
        favorites.remove(job_id)
    else:
        favorites.insert(0, job_id)

    return save_json(
        FAVORITES_PATH,
        favorites,
    )


# ---------------------------------------------------------
# ЧТЕНИЕ ЗАГРУЖЕННЫХ ФАЙЛОВ
# ---------------------------------------------------------

def read_plain_text_file(
    file_name: str,
    file_bytes: bytes,
) -> str:
    extension = Path(file_name).suffix.lower()

    if extension == ".csv":
        decoded = file_bytes.decode(
            "utf-8",
            errors="replace",
        )

        reader = csv.reader(
            io.StringIO(decoded)
        )

        rows = [
            " | ".join(row)
            for row in reader
        ]

        return "\n".join(rows)

    if extension == ".json":
        decoded = file_bytes.decode(
            "utf-8",
            errors="replace",
        )

        try:
            parsed = json.loads(decoded)

            return json.dumps(
                parsed,
                ensure_ascii=False,
                indent=2,
            )

        except json.JSONDecodeError:
            return decoded

    return file_bytes.decode(
        "utf-8",
        errors="replace",
    )


def read_docx_file(file_bytes: bytes) -> str:
    try:
        from docx import Document
    except Exception as exc:
        raise RuntimeError(
            "Для DOCX нужно добавить python-docx "
            "в requirements.txt."
        ) from exc

    document = Document(
        io.BytesIO(file_bytes)
    )

    paragraphs = [
        paragraph.text
        for paragraph in document.paragraphs
        if paragraph.text.strip()
    ]

    for table in document.tables:
        for row in table.rows:
            row_text = " | ".join(
                cell.text.strip()
                for cell in row.cells
            )

            if row_text.strip():
                paragraphs.append(row_text)

    return "\n".join(paragraphs)


def read_pdf_file(file_bytes: bytes) -> str:
    try:
        from pypdf import PdfReader
    except Exception as exc:
        raise RuntimeError(
            "Для PDF нужно добавить pypdf "
            "в requirements.txt."
        ) from exc

    reader = PdfReader(
        io.BytesIO(file_bytes)
    )

    pages: list[str] = []

    for page_number, page in enumerate(
        reader.pages,
        start=1,
    ):
        page_text = page.extract_text() or ""

        if page_text.strip():
            pages.append(
                f"\n--- Страница {page_number} ---\n"
                f"{page_text.strip()}"
            )

    if not pages:
        raise RuntimeError(
            "В PDF не удалось найти обычный текст. "
            "Возможно, это скан или изображение."
        )

    return "\n".join(pages)


def extract_uploaded_files(
    uploaded_files: list[Any] | None,
) -> tuple[str, list[str]]:
    if not uploaded_files:
        return "", []

    extracted_parts: list[str] = []
    messages: list[str] = []

    for uploaded_file in uploaded_files:
        file_name = str(
            uploaded_file.name
        )

        extension = Path(
            file_name
        ).suffix.lower()

        file_bytes = uploaded_file.getvalue()

        try:
            if extension in SUPPORTED_TEXT_EXTENSIONS:
                content = read_plain_text_file(
                    file_name,
                    file_bytes,
                )

            elif extension == ".docx":
                content = read_docx_file(
                    file_bytes
                )

            elif extension == ".pdf":
                content = read_pdf_file(
                    file_bytes
                )

            else:
                messages.append(
                    f"Файл «{file_name}» пропущен: "
                    "формат пока не поддерживается."
                )
                continue

            content = content.strip()

            if not content:
                messages.append(
                    f"В файле «{file_name}» "
                    "не найден текст."
                )
                continue

            extracted_parts.append(
                f"\n\n===== ФАЙЛ: {file_name} =====\n\n"
                f"{content[:MAX_FILE_TEXT_LENGTH]}"
            )

            messages.append(
                f"✅ Файл «{file_name}» прочитан."
            )

        except Exception as exc:
            messages.append(
                f"❌ Не удалось прочитать "
                f"«{file_name}»: {exc}"
            )

    return (
        "\n".join(extracted_parts).strip(),
        messages,
    )


# ---------------------------------------------------------
# ВЫБОР И ОТОБРАЖЕНИЕ ЗАКАЗА
# ---------------------------------------------------------

def select_job(
    jobs: list[dict[str, Any]],
    key: str,
    label: str = "Выбери заказ",
) -> dict[str, Any] | None:
    if not jobs:
        st.info(
            "Пока нет подходящих заказов."
        )
        return None

    selected_index = st.selectbox(
        label,
        options=list(range(len(jobs))),
        format_func=lambda index: job_option_label(
            jobs[index]
        ),
        key=key,
    )

    return jobs[selected_index]


def render_job_card(
    job: dict[str, Any],
    *,
    expanded_description: bool = True,
    card_key: str = "",
) -> None:
    job_id = str(
        job.get("id", "")
    )

    risk = str(
        job.get("risk", "medium")
    )

    risk_class = {
        "low": "risk-low",
        "medium": "risk-medium",
        "high": "risk-high",
    }.get(risk, "risk-medium")

    title = str(
        job.get("title", "Без названия")
    )

    company = str(
        job.get("company", "")
    ).strip()

    source = str(
        job.get("source", "Неизвестный источник")
    )

    opportunity_score = calculate_opportunity_score(
        job
    )

    simple_badge = (
        '<span class="badge simple-task">'
        '✨ Простое задание'
        '</span>'
        if is_simple_task(job)
        else ""
    )

    st.markdown(
        f"""
        <div class="job-card">
            <div class="job-title">{title}</div>
            <div class="job-meta">
                {company + " · " if company else ""}
                {source}
            </div>

            <span class="badge">
                ⭐ Выгодность: {opportunity_score}/100
            </span>

            <span class="badge {risk_class}">
                {risk_icon(job)} Риск: {risk_display(job)}
            </span>

            <span class="badge">
                💰 Бюджет: {budget_display(job)}
            </span>

            <span class="badge">
                🧩 {task_display(job)}
            </span>

            {simple_badge}
        </div>
        """,
        unsafe_allow_html=True,
    )

    c1, c2, c3, c4 = st.columns(4)

    c1.metric(
        "Сложность",
        estimated_difficulty(job),
    )

    c2.metric(
        "Примерное время",
        estimated_time(job),
    )

    c3.metric(
        "Рекомендованная цена",
        recommended_price(job),
    )

    c4.metric(
        "Язык",
        language_display(job),
    )

    if expanded_description:
        description = str(
            job.get("description", "")
        ).strip()

        if description:
            with st.expander(
                "Показать описание заказа",
                expanded=False,
            ):
                st.write(
                    description[:10_000]
                )

    matched_keywords = job.get(
        "matched_keywords",
        [],
    )

    if isinstance(matched_keywords, list) and matched_keywords:
        st.caption(
            "Подходящие слова: "
            + ", ".join(
                str(item)
                for item in matched_keywords[:15]
            )
        )

    risk_keywords = job.get(
        "risk_keywords",
        [],
    )

    if isinstance(risk_keywords, list) and risk_keywords:
        st.warning(
            "Подозрительные слова: "
            + ", ".join(
                str(item)
                for item in risk_keywords[:15]
            )
        )

    actions = st.columns([1, 1, 1, 2])

    favorite_text = (
        "💔 Убрать из избранного"
        if is_favorite(job_id)
        else "❤️ В избранное"
    )

    unique_key = (
        card_key
        or job_id
        or stable_id(
            source,
            str(job.get("link", "")),
            title,
        )
    )

    if actions[0].button(
        favorite_text,
        key=f"favorite_{unique_key}",
        use_container_width=True,
    ):
        if toggle_favorite(job_id):
            st.rerun()
        else:
            st.error(
                "Не удалось сохранить избранное."
            )

    link = str(
        job.get("link", "")
    ).strip()

    if link:
        actions[1].link_button(
            "🔗 Открыть оригинал",
            link,
            use_container_width=True,
        )

    if actions[2].button(
        "✉️ Выбрать для отклика",
        key=f"choose_proposal_{unique_key}",
        use_container_width=True,
    ):
        st.session_state[
            "selected_proposal_job_id"
        ] = job_id

        st.success(
            "Заказ выбран. Открой вкладку «Отклик»."
        )


# ---------------------------------------------------------
# ЗАГРУЗКА ДАННЫХ ПРИЛОЖЕНИЯ
# ---------------------------------------------------------

cfg = get_config()
profile = cfg.get("profile", {})

if not isinstance(profile, dict):
    profile = {}

repo_jobs = get_jobs()

manual_jobs = st.session_state.setdefault(
    "manual_jobs",
    [],
)

all_jobs = manual_jobs + repo_jobs

scan_status = load_json(
    STATUS_PATH,
    {},
)

favorites = get_favorites()
history = get_history()
# ---------------------------------------------------------
# ВЕРХНЯЯ ЧАСТЬ ИНТЕРФЕЙСА
# ---------------------------------------------------------

st.markdown(
    """
    <div class="hero-card">
        <div class="main-title">✨ Фриланс-Агент</div>
        <div class="main-subtitle">
            Поиск простых удалённых заданий, подготовка откликов,
            помощь с перепиской и выполнение текстовых заказов.
        </div>
    </div>
    """,
    unsafe_allow_html=True,
)


# ---------------------------------------------------------
# БОКОВАЯ ПАНЕЛЬ
# ---------------------------------------------------------

with st.sidebar:
    st.markdown("## ✨ Фриланс-Агент")

    if ai.available():
        st.success("OpenAI подключён")
    else:
        st.warning(
            "OpenAI не подключён. "
            "Доступны только базовые шаблоны."
        )

    st.divider()
    st.markdown("### 📊 Состояние")

    st.metric(
        "Всего заказов",
        len(repo_jobs),
    )

    st.metric(
        "Простых заданий",
        sum(
            is_simple_task(job)
            for job in repo_jobs
        ),
    )

    st.metric(
        "Низкий риск",
        sum(
            str(job.get("risk")) == "low"
            for job in repo_jobs
        ),
    )

    st.metric(
        "В избранном",
        len(favorites),
    )

    if scan_status:
        st.divider()
        st.markdown("### 🔄 Последнее сканирование")

        finished_at = str(
            scan_status.get(
                "finished_at",
                "Неизвестно",
            )
        )

        st.caption(
            f"Завершено: {finished_at}"
        )

        saved_total = scan_status.get(
            "saved_total",
            scan_status.get(
                "received_total",
                "—",
            ),
        )

        st.caption(
            f"Сохранено заданий: {saved_total}"
        )

        errors = scan_status.get(
            "errors",
            [],
        )

        if errors:
            st.warning(
                "Некоторые источники вернули ошибку."
            )

    st.divider()

    if st.button(
        "🔄 Обновить страницу",
        use_container_width=True,
    ):
        st.rerun()

    st.caption(
        "Приложение не отправляет заявки автоматически "
        "и не обходит ограничения площадок."
    )


# ---------------------------------------------------------
# ВКЛАДКИ
# ---------------------------------------------------------

(
    overview_tab,
    jobs_tab,
    favorites_tab,
    proposal_tab,
    chat_tab,
    work_tab,
    history_tab,
    diagnostics_tab,
) = st.tabs(
    [
        "🏠 Главная",
        "🔎 Заказы",
        "❤️ Избранное",
        "✉️ Отклик",
        "💬 Переписка",
        "💼 Работа",
        "📚 История",
        "⚙️ Диагностика",
    ]
)


# ---------------------------------------------------------
# ВКЛАДКА: ГЛАВНАЯ
# ---------------------------------------------------------

with overview_tab:
    st.subheader("Обзор")

    metric_1, metric_2, metric_3, metric_4 = st.columns(4)

    metric_1.metric(
        "Найдено заказов",
        len(repo_jobs),
    )

    metric_2.metric(
        "Подходящих",
        sum(
            calculate_opportunity_score(job) >= 65
            for job in repo_jobs
        ),
    )

    metric_3.metric(
        "Простых",
        sum(
            is_simple_task(job)
            for job in repo_jobs
        ),
    )

    metric_4.metric(
        "Высокий риск",
        sum(
            str(job.get("risk")) == "high"
            for job in repo_jobs
        ),
    )

    st.divider()

    left_column, right_column = st.columns(
        [1.6, 1]
    )

    with left_column:
        st.subheader("🔥 Лучшие найденные задания")

        best_jobs = sorted(
            repo_jobs,
            key=lambda job: (
                calculate_opportunity_score(job),
                int(job.get("score", 0)),
            ),
            reverse=True,
        )

        best_jobs = [
            job
            for job in best_jobs
            if str(job.get("risk")) != "high"
        ][:6]

        if not best_jobs:
            st.info(
                "Пока нет подходящих заданий. "
                "Дождись сканирования или добавь объявление вручную."
            )

        for index, job in enumerate(best_jobs):
            with st.expander(
                job_option_label(job),
                expanded=index == 0,
            ):
                render_job_card(
                    job,
                    expanded_description=True,
                    card_key=f"overview_{index}_{job.get('id', '')}",
                )

    with right_column:
        st.subheader("🎯 Что искать в первую очередь")

        st.markdown(
            """
            <div class="hero-card">
                <b>Лучшие простые категории:</b><br><br>
                🌍 перевод коротких текстов;<br>
                ✍️ проверка и исправление текста;<br>
                🎧 расшифровка аудио;<br>
                📋 ввод и перенос данных;<br>
                📄 PDF в Word;<br>
                🔎 простой поиск информации;<br>
                💬 письма и сообщения на немецком.
            </div>
            """,
            unsafe_allow_html=True,
        )

        st.subheader("🛡️ Основные правила безопасности")

        st.markdown(
            """
            <div class="hero-card">
                Не плати заказчику за получение работы.<br><br>
                Не покупай подарочные карты.<br><br>
                Не передавай доступ к банковскому счёту.<br><br>
                Не отправляй паспорт без понятной и законной причины.<br><br>
                Оплату лучше получать через безопасную систему площадки.
            </div>
            """,
            unsafe_allow_html=True,
        )

        st.subheader("📌 Быстрый статус")

        openai_status = (
            "Подключён"
            if ai.available()
            else "Не подключён"
        )

        st.write(
            f"**ИИ:** {openai_status}"
        )

        st.write(
            f"**Источников включено:** "
            f"{sum(bool(source.get('enabled')) for source in cfg.get('sources', []))}"
        )

        st.write(
            f"**Записей в истории:** {len(history)}"
        )


# ---------------------------------------------------------
# ВКЛАДКА: ЗАКАЗЫ
# ---------------------------------------------------------

with jobs_tab:
    st.subheader("🔎 Поиск и фильтрация заказов")

    with st.expander(
        "➕ Добавить объявление вручную",
        expanded=False,
    ):
        with st.form(
            "manual_job_form",
            clear_on_submit=True,
        ):
            manual_title = st.text_input(
                "Название задания",
                placeholder=(
                    "Например: Перевести короткий текст "
                    "с немецкого на русский"
                ),
            )

            manual_link = st.text_input(
                "Ссылка на объявление",
                placeholder="Необязательно",
            )

            manual_description = st.text_area(
                "Описание задания",
                height=220,
                placeholder=(
                    "Вставь сюда полное описание объявления."
                ),
            )

            manual_budget = st.text_input(
                "Бюджет",
                placeholder=(
                    "Например: $20 или $10–$30"
                ),
            )

            manual_submit = st.form_submit_button(
                "Добавить объявление",
                use_container_width=True,
            )

        if manual_submit:
            if not manual_title.strip():
                st.error(
                    "Нужно указать название задания."
                )

            elif not manual_description.strip():
                st.error(
                    "Нужно добавить описание задания."
                )

            else:
                combined_manual_text = (
                    f"{manual_title} "
                    f"{manual_description} "
                    f"{manual_budget}"
                )

                (
                    manual_score,
                    manual_matched,
                    manual_risks,
                    manual_risk,
                ) = score_job(
                    manual_title,
                    combined_manual_text,
                    cfg,
                )

                created_time = utc_now()

                manual_job = {
                    "id": stable_id(
                        "manual",
                        manual_link,
                        manual_title,
                    ),
                    "source": "Добавлено вручную",
                    "title": manual_title.strip(),
                    "link": manual_link.strip(),
                    "description": manual_description.strip(),
                    "published": created_time,
                    "scanned_at": created_time,
                    "company": "",
                    "location": "",
                    "category": "",
                    "language_hint": detect_language(
                        combined_manual_text
                    ),
                    "budget_text": (
                        manual_budget.strip()
                        or extract_budget(
                            combined_manual_text
                        )
                    ),
                    "score": manual_score,
                    "risk": manual_risk,
                    "task_type": classify_task_local(
                        {
                            "title": manual_title,
                            "description": manual_description,
                        }
                    ),
                    "is_simple": True,
                    "matched_keywords": manual_matched,
                    "risk_keywords": manual_risks,
                    "status": "new",
                }

                manual_jobs.insert(
                    0,
                    manual_job,
                )

                st.success(
                    "Объявление добавлено в текущую сессию."
                )

    st.divider()
    st.markdown("### Фильтры")

    filter_row_1 = st.columns(
        [1.1, 1.3, 1.1]
    )

    with filter_row_1[0]:
        minimum_score_filter = st.slider(
            "Минимальная выгодность",
            min_value=0,
            max_value=100,
            value=25,
            step=5,
        )

    with filter_row_1[1]:
        budget_filter = st.slider(
            "Бюджет задания, $",
            min_value=0,
            max_value=500,
            value=(5, 50),
            step=5,
            help=(
                "Показывает задания, бюджет которых "
                "пересекается с выбранным диапазоном."
            ),
        )

    with filter_row_1[2]:
        sorting_option = st.selectbox(
            "Сортировка",
            [
                "По выгодности",
                "По рейтингу сканера",
                "Сначала простые",
                "Сначала новые",
            ],
        )

    filter_row_2 = st.columns(
        [1.3, 1.1, 1.1]
    )

    with filter_row_2[0]:
        search_query = st.text_input(
            "Поиск по словам",
            placeholder=(
                "Например: перевод, немецкий, data entry"
            ),
        ).strip().lower()

    with filter_row_2[1]:
        risk_labels_selected = st.multiselect(
            "Допустимый риск",
            options=[
                "Низкий",
                "Средний",
                "Высокий",
            ],
            default=[
                "Низкий",
                "Средний",
            ],
        )

    with filter_row_2[2]:
        task_labels_selected = st.multiselect(
            "Тип задания",
            options=list(
                TASK_LABELS.values()
            ),
            default=[],
            placeholder="Все категории",
        )

    filter_row_3 = st.columns(3)

    with filter_row_3[0]:
        only_simple_filter = st.checkbox(
            "Только простые задания",
            value=True,
        )

    with filter_row_3[1]:
        include_without_budget_filter = st.checkbox(
            "Показывать без бюджета",
            value=True,
        )

    with filter_row_3[2]:
        only_favorites_filter = st.checkbox(
            "Только избранное",
            value=False,
        )

    selected_risk_values = {
        key
        for key, value in RISK_LABELS.items()
        if value in risk_labels_selected
    }

    filtered_jobs: list[dict[str, Any]] = []

    for job in all_jobs:
        opportunity_score = calculate_opportunity_score(
            job
        )

        if opportunity_score < minimum_score_filter:
            continue

        if (
            selected_risk_values
            and str(job.get("risk")) not in selected_risk_values
        ):
            continue

        if only_simple_filter and not is_simple_task(job):
            continue

        if only_favorites_filter and not is_favorite(
            str(job.get("id", ""))
        ):
            continue

        if task_labels_selected:
            current_task_label = task_display(job)

            if current_task_label not in task_labels_selected:
                continue

        if not budget_in_range(
            job,
            float(budget_filter[0]),
            float(budget_filter[1]),
            include_without_budget_filter,
        ):
            continue

        if search_query:
            searchable_text = job_full_text(job)

            if search_query not in searchable_text:
                continue

        filtered_jobs.append(job)

    if sorting_option == "По рейтингу сканера":
        filtered_jobs.sort(
            key=lambda job: int(
                job.get("score", 0)
            ),
            reverse=True,
        )

    elif sorting_option == "Сначала простые":
        filtered_jobs.sort(
            key=lambda job: (
                simple_task_score(job),
                calculate_opportunity_score(job),
            ),
            reverse=True,
        )

    elif sorting_option == "Сначала новые":
        filtered_jobs.sort(
            key=lambda job: str(
                job.get(
                    "published",
                    job.get("scanned_at", ""),
                )
            ),
            reverse=True,
        )

    else:
        filtered_jobs.sort(
            key=lambda job: calculate_opportunity_score(
                job
            ),
            reverse=True,
        )

    st.divider()

    result_header_1, result_header_2 = st.columns(
        [3, 1]
    )

    result_header_1.markdown(
        f"### Найдено: {len(filtered_jobs)}"
    )

    result_header_2.caption(
        f"Всего в базе: {len(all_jobs)}"
    )

    if not filtered_jobs:
        st.info(
            "По выбранным фильтрам ничего не найдено. "
            "Попробуй включить объявления без бюджета, "
            "расширить диапазон цены или отключить "
            "фильтр простых заданий."
        )

    jobs_to_show = filtered_jobs[:100]

    for index, job in enumerate(jobs_to_show):
        with st.expander(
            job_option_label(job),
            expanded=False,
        ):
            render_job_card(
                job,
                expanded_description=True,
                card_key=(
                    f"jobs_{index}_"
                    f"{job.get('id', '')}"
                ),
            )

    if len(filtered_jobs) > 100:
        st.caption(
            "Показаны первые 100 заданий. "
            "Используй фильтры, чтобы сузить результаты."
        )


# ---------------------------------------------------------
# ВКЛАДКА: ИЗБРАННОЕ
# ---------------------------------------------------------

with favorites_tab:
    st.subheader("❤️ Избранные задания")

    favorite_ids = set(
        get_favorites()
    )

    favorite_jobs = [
        job
        for job in all_jobs
        if str(job.get("id", "")) in favorite_ids
    ]

    favorite_jobs.sort(
        key=lambda job: calculate_opportunity_score(
            job
        ),
        reverse=True,
    )

    if not favorite_jobs:
        st.info(
            "В избранном пока ничего нет. "
            "Добавляй интересные задания кнопкой "
            "«В избранное»."
        )

    for index, job in enumerate(favorite_jobs):
        with st.expander(
            job_option_label(job),
            expanded=index == 0,
        ):
            render_job_card(
                job,
                expanded_description=True,
                card_key=(
                    f"favorites_{index}_"
                    f"{job.get('id', '')}"
                ),
            )
            # ---------------------------------------------------------
# УНИВЕРСАЛЬНЫЙ ВЫЗОВ OPENAI
# ---------------------------------------------------------

def get_openai_secret(
    secret_name: str,
    default: str = "",
) -> str:
    try:
        value = st.secrets.get(
            secret_name,
            default,
        )

        return str(value).strip()

    except Exception:
        return default


def extract_response_text(
    response: Any,
) -> str:
    direct_text = getattr(
        response,
        "output_text",
        None,
    )

    if direct_text:
        return str(direct_text).strip()

    output = getattr(
        response,
        "output",
        None,
    )

    if not output:
        return ""

    collected_parts: list[str] = []

    for output_item in output:
        content_items = getattr(
            output_item,
            "content",
            [],
        )

        for content_item in content_items:
            text_value = getattr(
                content_item,
                "text",
                None,
            )

            if text_value:
                collected_parts.append(
                    str(text_value)
                )

    return "\n".join(
        collected_parts
    ).strip()


def call_openai(
    instructions: str,
    user_content: str,
    *,
    max_output_tokens: int = 2200,
) -> str:
    api_key = get_openai_secret(
        "OPENAI_API_KEY"
    )

    model = get_openai_secret(
        "OPENAI_MODEL",
        "gpt-5-mini",
    )

    if not api_key:
        raise RuntimeError(
            "В Streamlit Secrets отсутствует "
            "OPENAI_API_KEY."
        )

    try:
        from openai import OpenAI
    except Exception as exc:
        raise RuntimeError(
            "Библиотека OpenAI не установлена. "
            "Добавь openai>=1.0.0 в requirements.txt."
        ) from exc

    client = OpenAI(
        api_key=api_key
    )

    response = client.responses.create(
        model=model,
        instructions=instructions,
        input=user_content,
        max_output_tokens=max_output_tokens,
    )

    result = extract_response_text(
        response
    )

    if not result:
        raise RuntimeError(
            "OpenAI вернул пустой результат."
        )

    return result


def ai_action(
    instructions: str,
    user_content: str,
    *,
    spinner_text: str,
    max_output_tokens: int = 2200,
) -> str | None:
    if not ai.available():
        st.error(
            "OpenAI не подключён. Проверь "
            "OPENAI_API_KEY в Streamlit Secrets."
        )
        return None

    try:
        with st.spinner(
            spinner_text
        ):
            return call_openai(
                instructions,
                user_content,
                max_output_tokens=max_output_tokens,
            )

    except Exception as exc:
        st.error(
            f"Не удалось выполнить запрос: {exc}"
        )
        return None


def selected_job_by_session(
    jobs: list[dict[str, Any]],
) -> dict[str, Any] | None:
    selected_job_id = str(
        st.session_state.get(
            "selected_proposal_job_id",
            "",
        )
    )

    if not selected_job_id:
        return None

    for job in jobs:
        if str(job.get("id", "")) == selected_job_id:
            return job

    return None


def proposal_fallback(
    job: dict[str, Any],
    style: str,
    language: str,
) -> str:
    title = str(
        job.get(
            "title",
            "your project",
        )
    )

    task_type = task_display(job)

    language_names = {
        "English": "English",
        "German": "German",
        "Russian": "Russian",
        "Ukrainian": "Ukrainian",
    }

    selected_language = language_names.get(
        language,
        "English",
    )

    if selected_language == "German":
        return (
            f"Guten Tag,\n\n"
            f"ich habe Ihre Aufgabe „{title}“ gelesen und "
            f"kann Sie dabei zuverlässig unterstützen. "
            f"Ich habe Erfahrung im Bereich {task_type.lower()} "
            f"und arbeite sorgfältig, verständlich und termingerecht.\n\n"
            f"Ich kann kurzfristig beginnen. Vor dem Start würde ich "
            f"gern den genauen Umfang, das gewünschte Format und die "
            f"Abgabefrist bestätigen.\n\n"
            f"Freundliche Grüße\n"
            f"{profile.get('display_name', 'Olha')}"
        )

    if selected_language == "Russian":
        return (
            f"Здравствуйте!\n\n"
            f"Я ознакомилась с заданием «{title}» и могу аккуратно "
            f"выполнить эту работу. У меня есть опыт в категории "
            f"«{task_type}», я внимательно соблюдаю инструкции "
            f"и сроки.\n\n"
            f"Могу начать в ближайшее время. Перед началом предлагаю "
            f"уточнить объём, формат результата и срок сдачи.\n\n"
            f"С уважением,\n"
            f"{profile.get('display_name', 'Olha')}"
        )

    if selected_language == "Ukrainian":
        return (
            f"Вітаю!\n\n"
            f"Я ознайомилася із завданням «{title}» і можу акуратно "
            f"виконати цю роботу. Маю досвід у категорії "
            f"«{task_type}», уважно дотримуюся інструкцій "
            f"і термінів.\n\n"
            f"Можу розпочати найближчим часом. Перед початком варто "
            f"уточнити обсяг, формат результату та кінцевий термін.\n\n"
            f"З повагою,\n"
            f"{profile.get('display_name', 'Olha')}"
        )

    return (
        f"Hello,\n\n"
        f"I reviewed your project “{title}” and I can help you "
        f"complete it carefully and on time. I have relevant "
        f"experience with {task_type.lower()} and I pay close "
        f"attention to instructions, formatting, and quality.\n\n"
        f"I am available to start soon. Before beginning, I would "
        f"like to confirm the exact scope, required format, and deadline.\n\n"
        f"Best regards,\n"
        f"{profile.get('display_name', 'Olha')}"
    )


# ---------------------------------------------------------
# ВКЛАДКА: ОТКЛИК
# ---------------------------------------------------------

with proposal_tab:
    st.subheader("✉️ Создание отклика")

    st.info(
        "Приложение создаёт текст отклика, но не отправляет "
        "его на площадку автоматически."
    )

    session_selected_job = selected_job_by_session(
        all_jobs
    )

    default_job_index = 0

    if session_selected_job:
        for index, job in enumerate(all_jobs):
            if str(job.get("id", "")) == str(
                session_selected_job.get("id", "")
            ):
                default_job_index = index
                break

    if all_jobs:
        selected_proposal_index = st.selectbox(
            "Заказ для отклика",
            options=list(range(len(all_jobs))),
            index=min(
                default_job_index,
                len(all_jobs) - 1,
            ),
            format_func=lambda index: job_option_label(
                all_jobs[index]
            ),
            key="proposal_job_selector",
        )

        proposal_job = all_jobs[
            selected_proposal_index
        ]

        st.session_state[
            "selected_proposal_job_id"
        ] = str(
            proposal_job.get("id", "")
        )

        render_job_card(
            proposal_job,
            expanded_description=False,
            card_key="proposal_selected_job",
        )

        st.divider()

        proposal_settings_1 = st.columns(3)

        with proposal_settings_1[0]:
            proposal_language = st.selectbox(
                "Язык отклика",
                [
                    "English",
                    "German",
                    "Russian",
                    "Ukrainian",
                ],
                index=0,
            )

        with proposal_settings_1[1]:
            proposal_style = st.selectbox(
                "Стиль",
                [
                    "Короткий",
                    "Профессиональный",
                    "Дружелюбный",
                    "Уверенный",
                ],
                index=1,
            )

        with proposal_settings_1[2]:
            proposal_length = st.selectbox(
                "Длина",
                [
                    "Очень короткий",
                    "Обычный",
                    "Подробный",
                ],
                index=1,
            )

        proposal_settings_2 = st.columns(2)

        with proposal_settings_2[0]:
            proposed_price = st.text_input(
                "Предлагаемая цена",
                value=recommended_price(
                    proposal_job
                ),
                help=(
                    "Укажи только число или сумму так, "
                    "как она должна выглядеть в отклике."
                ),
            )

        with proposal_settings_2[1]:
            proposed_deadline = st.text_input(
                "Срок выполнения",
                value=estimated_time(
                    proposal_job
                ),
            )

        relevant_experience = st.text_area(
            "Опыт или важная информация о себе",
            value=(
                "I speak Ukrainian, Russian, German and English. "
                "I work carefully, follow instructions and communicate clearly."
            ),
            height=110,
        )

        extra_proposal_instruction = st.text_area(
            "Дополнительные пожелания",
            placeholder=(
                "Например: не писать слишком официально, "
                "не обещать невозможного, задать один важный вопрос."
            ),
            height=100,
        )

        proposal_buttons = st.columns(
            [1, 1, 1]
        )

        generate_proposal_clicked = proposal_buttons[0].button(
            "✨ Создать отклик",
            type="primary",
            use_container_width=True,
        )

        fallback_proposal_clicked = proposal_buttons[1].button(
            "📝 Создать без ИИ",
            use_container_width=True,
        )

        clear_proposal_clicked = proposal_buttons[2].button(
            "🗑️ Очистить",
            use_container_width=True,
        )

        if clear_proposal_clicked:
            st.session_state.pop(
                "generated_proposal",
                None,
            )
            st.rerun()

        if fallback_proposal_clicked:
            st.session_state[
                "generated_proposal"
            ] = proposal_fallback(
                proposal_job,
                proposal_style,
                proposal_language,
            )

        if generate_proposal_clicked:
            proposal_prompt = f"""
JOB TITLE:
{proposal_job.get("title", "")}

JOB DESCRIPTION:
{proposal_job.get("description", "")[:12000]}

JOB CATEGORY:
{task_display(proposal_job)}

JOB BUDGET:
{budget_display(proposal_job)}

PROPOSED PRICE:
{proposed_price}

PROPOSED DELIVERY TIME:
{proposed_deadline}

FREELANCER NAME:
{profile.get("display_name", "Olha")}

FREELANCER LANGUAGES:
{", ".join(profile.get("languages", []))}

FREELANCER SERVICES:
{", ".join(profile.get("services", []))}

RELEVANT EXPERIENCE:
{relevant_experience}

REQUESTED STYLE:
{proposal_style}

REQUESTED LENGTH:
{proposal_length}

OUTPUT LANGUAGE:
{proposal_language}

EXTRA INSTRUCTIONS:
{extra_proposal_instruction}
"""

            generated_proposal = ai_action(
                instructions=(
                    "Create a realistic freelance proposal for the job. "
                    "Do not invent qualifications, previous clients, ratings, "
                    "certificates, or completed projects. Do not guarantee "
                    "results that cannot be guaranteed. Make the proposal "
                    "specific to the job description. Mention price and "
                    "delivery time only when they are useful. Include no more "
                    "than two relevant questions. Output only the finished "
                    "proposal without analysis, headings, or quotation marks."
                ),
                user_content=proposal_prompt,
                spinner_text="Создаю отклик…",
                max_output_tokens=1400,
            )

            if generated_proposal:
                st.session_state[
                    "generated_proposal"
                ] = generated_proposal

                add_history_entry(
                    "proposal",
                    str(
                        proposal_job.get(
                            "title",
                            "Отклик",
                        )
                    ),
                    generated_proposal,
                    str(
                        proposal_job.get(
                            "id",
                            "",
                        )
                    ),
                )

        generated_proposal_value = str(
            st.session_state.get(
                "generated_proposal",
                "",
            )
        )

        if generated_proposal_value:
            st.divider()
            st.markdown("### Готовый отклик")

            edited_proposal = st.text_area(
                "Текст можно отредактировать",
                value=generated_proposal_value,
                height=320,
                key="proposal_editor",
            )

            proposal_result_buttons = st.columns(3)

            if proposal_result_buttons[0].button(
                "✨ Улучшить текст",
                use_container_width=True,
            ):
                improved_proposal = ai_action(
                    instructions=(
                        "Improve the freelance proposal while preserving "
                        "all true facts. Make it specific, natural, concise, "
                        "and persuasive. Remove repetition and generic filler. "
                        "Do not invent experience or achievements. Output only "
                        "the improved proposal."
                    ),
                    user_content=edited_proposal,
                    spinner_text="Улучшаю отклик…",
                    max_output_tokens=1400,
                )

                if improved_proposal:
                    st.session_state[
                        "generated_proposal"
                    ] = improved_proposal

                    add_history_entry(
                        "proposal_improved",
                        str(
                            proposal_job.get(
                                "title",
                                "Улучшенный отклик",
                            )
                        ),
                        improved_proposal,
                        str(
                            proposal_job.get(
                                "id",
                                "",
                            )
                        ),
                    )

                    st.rerun()

            proposal_result_buttons[1].download_button(
                "⬇️ Скачать TXT",
                data=edited_proposal,
                file_name="proposal.txt",
                mime="text/plain",
                use_container_width=True,
            )

            if proposal_result_buttons[2].button(
                "💾 Сохранить в историю",
                use_container_width=True,
            ):
                if add_history_entry(
                    "proposal_saved",
                    str(
                        proposal_job.get(
                            "title",
                            "Отклик",
                        )
                    ),
                    edited_proposal,
                    str(
                        proposal_job.get(
                            "id",
                            "",
                        )
                    ),
                ):
                    st.success(
                        "Отклик сохранён в историю."
                    )
                else:
                    st.error(
                        "Не удалось сохранить историю."
                    )

            st.caption(
                "Для копирования нажми внутри поля, затем "
                "используй стандартную команду копирования."
            )

    else:
        st.info(
            "В базе пока нет заказов. Добавь объявление "
            "вручную во вкладке «Заказы»."
        )


# ---------------------------------------------------------
# ВКЛАДКА: ПЕРЕПИСКА
# ---------------------------------------------------------

with chat_tab:
    st.subheader("💬 Помощник для переписки")

    chat_columns = st.columns(
        [1.2, 1]
    )

    with chat_columns[0]:
        client_message = st.text_area(
            "Сообщение клиента",
            height=240,
            placeholder=(
                "Вставь сюда сообщение, которое прислал клиент."
            ),
        )

        reply_context = st.text_area(
            "Что ты хочешь ответить",
            height=140,
            placeholder=(
                "Например: я могу закончить завтра вечером, "
                "но мне нужен исходный файл в DOCX."
            ),
        )

    with chat_columns[1]:
        reply_language = st.selectbox(
            "Язык ответа",
            [
                "English",
                "German",
                "Russian",
                "Ukrainian",
            ],
            key="chat_reply_language",
        )

        reply_tone = st.selectbox(
            "Тон ответа",
            [
                "Вежливый",
                "Дружелюбный",
                "Профессиональный",
                "Короткий и прямой",
            ],
        )

        reply_goal = st.selectbox(
            "Цель",
            [
                "Ответить на сообщение",
                "Уточнить детали",
                "Согласовать цену",
                "Согласовать срок",
                "Вежливо отказаться",
                "Сообщить о готовности работы",
            ],
        )

        include_translation = st.checkbox(
            "Показать перевод сообщения клиента на русский",
            value=True,
        )

    chat_action_columns = st.columns(3)

    generate_reply_clicked = chat_action_columns[0].button(
        "✨ Создать ответ",
        type="primary",
        use_container_width=True,
    )

    translate_only_clicked = chat_action_columns[1].button(
        "🌍 Только перевести",
        use_container_width=True,
    )

    clear_chat_clicked = chat_action_columns[2].button(
        "🗑️ Очистить результат",
        use_container_width=True,
    )

    if clear_chat_clicked:
        st.session_state.pop(
            "chat_result",
            None,
        )
        st.session_state.pop(
            "chat_translation",
            None,
        )
        st.rerun()

    if translate_only_clicked:
        if not client_message.strip():
            st.warning(
                "Сначала вставь сообщение клиента."
            )
        else:
            translated_message = ai_action(
                instructions=(
                    "Translate the supplied message into Russian. "
                    "Preserve meaning, prices, dates, names, and formatting. "
                    "Do not add explanations. Output only the translation."
                ),
                user_content=client_message,
                spinner_text="Перевожу сообщение…",
                max_output_tokens=1600,
            )

            if translated_message:
                st.session_state[
                    "chat_translation"
                ] = translated_message

    if generate_reply_clicked:
        if not client_message.strip():
            st.warning(
                "Сначала вставь сообщение клиента."
            )

        elif not reply_context.strip():
            st.warning(
                "Коротко напиши, что именно ты хочешь ответить."
            )

        else:
            chat_prompt = f"""
CLIENT MESSAGE:
{client_message}

WHAT I WANT TO SAY:
{reply_context}

REPLY LANGUAGE:
{reply_language}

TONE:
{reply_tone}

GOAL:
{reply_goal}

FREELANCER NAME:
{profile.get("display_name", "Olha")}
"""

            generated_reply = ai_action(
                instructions=(
                    "Write a natural message to a freelance client. "
                    "Use only the facts provided by the user. Do not invent "
                    "completed work, deadlines, prices, attachments, or "
                    "agreements. Keep the reply clear and appropriate for "
                    "professional communication. Output only the message."
                ),
                user_content=chat_prompt,
                spinner_text="Готовлю ответ клиенту…",
                max_output_tokens=1200,
            )

            if generated_reply:
                st.session_state[
                    "chat_result"
                ] = generated_reply

                add_history_entry(
                    "client_reply",
                    "Ответ клиенту",
                    generated_reply,
                )

            if include_translation:
                translated_message = ai_action(
                    instructions=(
                        "Translate the supplied client message into Russian. "
                        "Preserve its meaning and details. Output only the "
                        "translation."
                    ),
                    user_content=client_message,
                    spinner_text="Перевожу сообщение клиента…",
                    max_output_tokens=1400,
                )

                if translated_message:
                    st.session_state[
                        "chat_translation"
                    ] = translated_message

    saved_translation = str(
        st.session_state.get(
            "chat_translation",
            "",
        )
    )

    saved_reply = str(
        st.session_state.get(
            "chat_result",
            "",
        )
    )

    if saved_translation:
        st.divider()
        st.markdown("### Перевод сообщения клиента")

        st.text_area(
            "Перевод",
            value=saved_translation,
            height=180,
            key="chat_translation_editor",
        )

    if saved_reply:
        st.markdown("### Готовый ответ")

        edited_reply = st.text_area(
            "Ответ можно отредактировать",
            value=saved_reply,
            height=220,
            key="chat_reply_editor",
        )

        reply_download_columns = st.columns(2)

        reply_download_columns[0].download_button(
            "⬇️ Скачать ответ",
            data=edited_reply,
            file_name="client_reply.txt",
            mime="text/plain",
            use_container_width=True,
        )

        if reply_download_columns[1].button(
            "💾 Сохранить ответ",
            use_container_width=True,
        ):
            if add_history_entry(
                "client_reply_saved",
                "Сохранённый ответ клиенту",
                edited_reply,
            ):
                st.success(
                    "Ответ сохранён в историю."
                )


# ---------------------------------------------------------
# ВКЛАДКА: РАБОТА
# ---------------------------------------------------------

with work_tab:
    st.subheader("💼 Выполнение текстового задания")

    st.warning(
        "Перед отправкой клиенту обязательно проверь результат. "
        "ИИ может неправильно понять инструкцию, числа или термины."
    )

    work_mode = st.selectbox(
        "Что нужно сделать",
        [
            "Выполнить задание по инструкции",
            "Перевести текст",
            "Проверить и исправить текст",
            "Сделать краткое содержание",
            "Переписать более профессионально",
            "Извлечь данные в таблицу",
            "Проверить готовый результат",
        ],
    )

    uploaded_files = st.file_uploader(
        "Загрузить файлы",
        type=[
            "txt",
            "md",
            "csv",
            "json",
            "docx",
            "pdf",
        ],
        accept_multiple_files=True,
        help=(
            "Поддерживаются TXT, MD, CSV, JSON, DOCX и PDF "
            "с обычным текстовым слоем."
        ),
    )

    extracted_file_text, file_messages = extract_uploaded_files(
        uploaded_files
    )

    for message in file_messages:
        if message.startswith("✅"):
            st.success(message)
        elif message.startswith("❌"):
            st.error(message)
        else:
            st.warning(message)

    work_instruction = st.text_area(
        "Инструкция клиента",
        height=180,
        placeholder=(
            "Вставь полное задание клиента. Например: "
            "перевести текст на немецкий, сохранить структуру "
            "и не переводить названия брендов."
        ),
    )

    manual_work_text = st.text_area(
        "Исходный текст",
        height=260,
        placeholder=(
            "Можно вставить текст вручную или загрузить файл выше."
        ),
    )

    combined_work_input = "\n\n".join(
        part
        for part in [
            manual_work_text.strip(),
            extracted_file_text.strip(),
        ]
        if part
    )

    work_options = st.columns(3)

    with work_options[0]:
        work_output_language = st.selectbox(
            "Язык результата",
            [
                "Согласно инструкции",
                "English",
                "German",
                "Russian",
                "Ukrainian",
            ],
        )

    with work_options[1]:
        work_quality = st.selectbox(
            "Уровень обработки",
            [
                "Обычный",
                "Тщательный",
                "Максимально тщательный",
            ],
            index=1,
        )

    with work_options[2]:
        preserve_formatting = st.checkbox(
            "Сохранять структуру",
            value=True,
        )

    additional_work_rules = st.text_area(
        "Дополнительные правила",
        placeholder=(
            "Например: не изменять числа, названия компаний, "
            "ссылки и медицинские термины."
        ),
        height=100,
    )

    work_buttons = st.columns(3)

    perform_work_clicked = work_buttons[0].button(
        "🤖 Выполнить",
        type="primary",
        use_container_width=True,
    )

    check_work_clicked = work_buttons[1].button(
        "🔍 Проверить результат",
        use_container_width=True,
    )

    clear_work_clicked = work_buttons[2].button(
        "🗑️ Очистить результат",
        use_container_width=True,
    )

    if clear_work_clicked:
        st.session_state.pop(
            "work_result",
            None,
        )
        st.session_state.pop(
            "quality_report",
            None,
        )
        st.rerun()

    if perform_work_clicked:
        if not work_instruction.strip():
            st.warning(
                "Добавь инструкцию клиента."
            )

        elif not combined_work_input.strip():
            st.warning(
                "Добавь исходный текст или загрузи файл."
            )

        else:
            work_prompt = f"""
TASK TYPE:
{work_mode}

CLIENT INSTRUCTIONS:
{work_instruction}

REQUESTED OUTPUT LANGUAGE:
{work_output_language}

QUALITY LEVEL:
{work_quality}

PRESERVE STRUCTURE:
{preserve_formatting}

ADDITIONAL RULES:
{additional_work_rules}

SOURCE CONTENT:
{combined_work_input[:85000]}
"""

            completed_work = ai_action(
                instructions=(
                    "Complete the supplied text-based freelance task. "
                    "Follow the client's instructions exactly. Do not invent "
                    "missing information. Preserve numbers, names, links, "
                    "technical terms, and document structure unless the user "
                    "explicitly asks to change them. If essential information "
                    "is missing, clearly mark what could not be completed "
                    "instead of guessing. Output only the finished deliverable."
                ),
                user_content=work_prompt,
                spinner_text="Выполняю задание…",
                max_output_tokens=5000,
            )

            if completed_work:
                st.session_state[
                    "work_result"
                ] = completed_work

                add_history_entry(
                    "completed_work",
                    work_mode,
                    completed_work,
                )

    current_work_result = str(
        st.session_state.get(
            "work_result",
            "",
        )
    )

    if check_work_clicked:
        if not current_work_result.strip():
            st.warning(
                "Сначала выполни задание или вставь результат."
            )
        else:
            quality_prompt = f"""
CLIENT INSTRUCTIONS:
{work_instruction}

SOURCE CONTENT:
{combined_work_input[:45000]}

RESULT TO CHECK:
{current_work_result[:45000]}
"""

            quality_report = ai_action(
                instructions=(
                    "Act as a quality-control reviewer. Compare the result "
                    "against the client instructions and source content. "
                    "Check for missing information, invented facts, changed "
                    "numbers, incorrect names, translation errors, formatting "
                    "problems, and unclear wording. Respond in Russian. "
                    "Provide: 1) overall verdict, 2) specific issues, "
                    "3) corrected version only when corrections are needed. "
                    "Be precise and do not claim certainty when source "
                    "information is insufficient."
                ),
                user_content=quality_prompt,
                spinner_text="Проверяю качество…",
                max_output_tokens=4000,
            )

            if quality_report:
                st.session_state[
                    "quality_report"
                ] = quality_report

                add_history_entry(
                    "quality_check",
                    "Проверка выполненной работы",
                    quality_report,
                )

    current_work_result = str(
        st.session_state.get(
            "work_result",
            "",
        )
    )

    if current_work_result:
        st.divider()
        st.markdown("### Готовый результат")

        edited_work_result = st.text_area(
            "Проверь и при необходимости отредактируй",
            value=current_work_result,
            height=480,
            key="work_result_editor",
        )

        work_result_actions = st.columns(3)

        if work_result_actions[0].button(
            "✨ Исправить и улучшить",
            use_container_width=True,
        ):
            improved_work = ai_action(
                instructions=(
                    "Improve and correct the supplied work result according "
                    "to the client instructions. Preserve correct content, "
                    "facts, numbers, names, and formatting. Fix only genuine "
                    "issues. Do not add invented information. Output only the "
                    "final corrected deliverable."
                ),
                user_content=(
                    f"CLIENT INSTRUCTIONS:\n{work_instruction}\n\n"
                    f"RESULT:\n{edited_work_result}"
                ),
                spinner_text="Исправляю результат…",
                max_output_tokens=5000,
            )

            if improved_work:
                st.session_state[
                    "work_result"
                ] = improved_work
                st.rerun()

        work_result_actions[1].download_button(
            "⬇️ Скачать TXT",
            data=edited_work_result,
            file_name="completed_work.txt",
            mime="text/plain",
            use_container_width=True,
        )

        if work_result_actions[2].button(
            "💾 Сохранить в историю",
            use_container_width=True,
        ):
            if add_history_entry(
                "completed_work_saved",
                work_mode,
                edited_work_result,
            ):
                st.success(
                    "Результат сохранён."
                )

    quality_report_value = str(
        st.session_state.get(
            "quality_report",
            "",
        )
    )

    if quality_report_value:
        st.divider()
        st.markdown("### 🔍 Отчёт о проверке")

        st.text_area(
            "Результат проверки",
            value=quality_report_value,
            height=380,
            key="quality_report_editor",
        )
        # ---------------------------------------------------------
# ВКЛАДКА: ИСТОРИЯ
# ---------------------------------------------------------

with history_tab:
    st.subheader("📚 История работы")

    current_history = get_history()

    if not current_history:
        st.info(
            "История пока пустая. Здесь появятся созданные отклики, "
            "ответы клиентам, выполненные задания и проверки качества."
        )

    else:
        history_statistics = st.columns(4)

        history_statistics[0].metric(
            "Всего записей",
            len(current_history),
        )

        history_statistics[1].metric(
            "Отклики",
            sum(
                str(item.get("type", "")).startswith("proposal")
                for item in current_history
            ),
        )

        history_statistics[2].metric(
            "Выполненные работы",
            sum(
                str(item.get("type", "")).startswith("completed_work")
                for item in current_history
            ),
        )

        history_statistics[3].metric(
            "Ответы клиентам",
            sum(
                str(item.get("type", "")).startswith("client_reply")
                for item in current_history
            ),
        )

        st.divider()

        history_filter_columns = st.columns(
            [1.2, 1.2, 1]
        )

        with history_filter_columns[0]:
            history_type_filter = st.selectbox(
                "Тип записи",
                [
                    "Все",
                    "Отклики",
                    "Ответы клиентам",
                    "Выполненные работы",
                    "Проверки качества",
                    "Другое",
                ],
                key="history_type_filter",
            )

        with history_filter_columns[1]:
            history_search = st.text_input(
                "Поиск в истории",
                placeholder="Название или текст",
                key="history_search",
            ).strip().lower()

        with history_filter_columns[2]:
            history_limit = st.selectbox(
                "Показывать записей",
                [
                    20,
                    50,
                    100,
                    300,
                ],
                index=1,
                key="history_limit",
            )

        def history_category(
            history_item: dict[str, Any],
        ) -> str:
            item_type = str(
                history_item.get("type", "")
            )

            if item_type.startswith("proposal"):
                return "Отклики"

            if item_type.startswith("client_reply"):
                return "Ответы клиентам"

            if item_type.startswith("completed_work"):
                return "Выполненные работы"

            if item_type.startswith("quality"):
                return "Проверки качества"

            return "Другое"

        filtered_history: list[dict[str, Any]] = []

        for history_item in current_history:
            item_category = history_category(
                history_item
            )

            if (
                history_type_filter != "Все"
                and item_category != history_type_filter
            ):
                continue

            if history_search:
                history_searchable_text = (
                    f"{history_item.get('title', '')} "
                    f"{history_item.get('content', '')} "
                    f"{history_item.get('type', '')}"
                ).lower()

                if history_search not in history_searchable_text:
                    continue

            filtered_history.append(
                history_item
            )

        filtered_history = filtered_history[
            :history_limit
        ]

        st.caption(
            f"Показано записей: {len(filtered_history)}"
        )

        history_export_data = json.dumps(
            current_history,
            ensure_ascii=False,
            indent=2,
        )

        export_columns = st.columns(
            [1, 1, 2]
        )

        export_columns[0].download_button(
            "⬇️ Скачать историю JSON",
            data=history_export_data,
            file_name="freelance_agent_history.json",
            mime="application/json",
            use_container_width=True,
        )

        history_csv_buffer = io.StringIO()

        history_csv_writer = csv.writer(
            history_csv_buffer
        )

        history_csv_writer.writerow(
            [
                "Дата",
                "Тип",
                "Название",
                "Содержание",
                "ID заказа",
            ]
        )

        for history_item in current_history:
            history_csv_writer.writerow(
                [
                    history_item.get(
                        "created_at",
                        "",
                    ),
                    history_item.get(
                        "type",
                        "",
                    ),
                    history_item.get(
                        "title",
                        "",
                    ),
                    history_item.get(
                        "content",
                        "",
                    ),
                    history_item.get(
                        "job_id",
                        "",
                    ),
                ]
            )

        export_columns[1].download_button(
            "⬇️ Скачать историю CSV",
            data=history_csv_buffer.getvalue(),
            file_name="freelance_agent_history.csv",
            mime="text/csv",
            use_container_width=True,
        )

        if export_columns[2].button(
            "🗑️ Очистить всю историю",
            use_container_width=True,
        ):
            st.session_state[
                "confirm_clear_history"
            ] = True

        if st.session_state.get(
            "confirm_clear_history",
            False,
        ):
            st.warning(
                "Все записи истории будут удалены без возможности "
                "восстановления."
            )

            confirm_history_columns = st.columns(
                [1, 1, 2]
            )

            if confirm_history_columns[0].button(
                "Да, удалить",
                type="primary",
                use_container_width=True,
                key="confirm_clear_history_yes",
            ):
                if save_json(
                    HISTORY_PATH,
                    [],
                ):
                    st.session_state[
                        "confirm_clear_history"
                    ] = False

                    st.success(
                        "История очищена."
                    )

                    st.rerun()

                else:
                    st.error(
                        "Не удалось очистить историю."
                    )

            if confirm_history_columns[1].button(
                "Отмена",
                use_container_width=True,
                key="confirm_clear_history_no",
            ):
                st.session_state[
                    "confirm_clear_history"
                ] = False

                st.rerun()

        st.divider()

        history_type_labels = {
            "proposal": "✉️ Отклик",
            "proposal_saved": "💾 Сохранённый отклик",
            "proposal_improved": "✨ Улучшенный отклик",
            "client_reply": "💬 Ответ клиенту",
            "client_reply_saved": "💾 Сохранённый ответ",
            "completed_work": "💼 Выполненная работа",
            "completed_work_saved": "💾 Сохранённая работа",
            "quality_check": "🔍 Проверка качества",
        }

        for history_index, history_item in enumerate(
            filtered_history
        ):
            history_item_id = str(
                history_item.get(
                    "id",
                    history_index,
                )
            )

            history_item_type = str(
                history_item.get(
                    "type",
                    "other",
                )
            )

            history_item_title = str(
                history_item.get(
                    "title",
                    "Без названия",
                )
            )

            history_item_content = str(
                history_item.get(
                    "content",
                    "",
                )
            )

            history_item_created = str(
                history_item.get(
                    "created_at",
                    "",
                )
            )

            history_item_label = history_type_labels.get(
                history_item_type,
                "📄 Запись",
            )

            expander_title = (
                f"{history_item_label} · "
                f"{history_item_title[:90]}"
            )

            with st.expander(
                expander_title,
                expanded=False,
            ):
                st.caption(
                    f"Создано: {history_item_created or 'Неизвестно'}"
                )

                if history_item.get("job_id"):
                    st.caption(
                        f"ID заказа: {history_item.get('job_id')}"
                    )

                edited_history_content = st.text_area(
                    "Содержание",
                    value=history_item_content,
                    height=280,
                    key=(
                        f"history_content_"
                        f"{history_index}_"
                        f"{history_item_id}"
                    ),
                )

                history_item_actions = st.columns(
                    [1, 1, 1]
                )

                history_item_actions[0].download_button(
                    "⬇️ Скачать TXT",
                    data=edited_history_content,
                    file_name=(
                        f"history_{history_index + 1}.txt"
                    ),
                    mime="text/plain",
                    use_container_width=True,
                    key=(
                        f"history_download_"
                        f"{history_index}_"
                        f"{history_item_id}"
                    ),
                )

                if history_item_actions[1].button(
                    "💾 Сохранить изменения",
                    use_container_width=True,
                    key=(
                        f"history_save_"
                        f"{history_index}_"
                        f"{history_item_id}"
                    ),
                ):
                    complete_history = get_history()
                    history_updated = False

                    for complete_item in complete_history:
                        if str(
                            complete_item.get(
                                "id",
                                "",
                            )
                        ) == history_item_id:
                            complete_item[
                                "content"
                            ] = edited_history_content

                            complete_item[
                                "updated_at"
                            ] = utc_now()

                            history_updated = True
                            break

                    if history_updated and save_json(
                        HISTORY_PATH,
                        complete_history,
                    ):
                        st.success(
                            "Изменения сохранены."
                        )

                    else:
                        st.error(
                            "Не удалось сохранить изменения."
                        )

                if history_item_actions[2].button(
                    "🗑️ Удалить запись",
                    use_container_width=True,
                    key=(
                        f"history_delete_"
                        f"{history_index}_"
                        f"{history_item_id}"
                    ),
                ):
                    complete_history = get_history()

                    new_history = [
                        item
                        for item in complete_history
                        if str(
                            item.get(
                                "id",
                                "",
                            )
                        ) != history_item_id
                    ]

                    if save_json(
                        HISTORY_PATH,
                        new_history,
                    ):
                        st.success(
                            "Запись удалена."
                        )

                        st.rerun()

                    else:
                        st.error(
                            "Не удалось удалить запись."
                        )


# ---------------------------------------------------------
# ВКЛАДКА: ДИАГНОСТИКА
# ---------------------------------------------------------

with diagnostics_tab:
    st.subheader("⚙️ Диагностика приложения")

    st.caption(
        "Здесь можно проверить подключение OpenAI, файлы проекта, "
        "источники заданий и основные настройки."
    )

    diagnostic_openai_key = get_openai_secret(
        "OPENAI_API_KEY"
    )

    diagnostic_openai_model = get_openai_secret(
        "OPENAI_MODEL",
        "gpt-5-mini",
    )

    diagnostic_columns = st.columns(4)

    diagnostic_columns[0].metric(
        "OpenAI API",
        (
            "Подключён"
            if ai.available()
            else "Не подключён"
        ),
    )

    diagnostic_columns[1].metric(
        "API-ключ",
        (
            "Найден"
            if diagnostic_openai_key
            else "Не найден"
        ),
    )

    diagnostic_columns[2].metric(
        "Модель",
        diagnostic_openai_model,
    )

    diagnostic_columns[3].metric(
        "Заданий в базе",
        len(repo_jobs),
    )

    st.divider()

    diagnostic_section_1, diagnostic_section_2 = st.columns(
        2
    )

    with diagnostic_section_1:
        st.markdown("### 🤖 OpenAI")

        if ai.available():
            st.success(
                "Модуль ИИ сообщает, что OpenAI доступен."
            )
        else:
            st.error(
                "OpenAI недоступен. Проверь Streamlit Secrets "
                "и файл requirements.txt."
            )

        if diagnostic_openai_key:
            masked_key = (
                diagnostic_openai_key[:7]
                + "..."
                + diagnostic_openai_key[-4:]
                if len(diagnostic_openai_key) >= 15
                else "Ключ найден"
            )

            st.code(
                f"OPENAI_API_KEY = {masked_key}"
            )
        else:
            st.code(
                "OPENAI_API_KEY = не найден"
            )

        st.code(
            f"OPENAI_MODEL = {diagnostic_openai_model}"
        )

        if st.button(
            "🧪 Проверить запрос к OpenAI",
            type="primary",
            use_container_width=True,
            key="test_openai_connection",
        ):
            test_result = ai_action(
                instructions=(
                    "Answer in Russian with one short sentence. "
                    "Confirm that the API request works. "
                    "Do not add headings or explanations."
                ),
                user_content=(
                    "Проверка подключения приложения к OpenAI."
                ),
                spinner_text="Проверяю OpenAI…",
                max_output_tokens=100,
            )

            if test_result:
                st.success(
                    "Запрос выполнен успешно."
                )

                st.code(
                    test_result
                )

    with diagnostic_section_2:
        st.markdown("### 📁 Файлы проекта")

        required_files = {
            "app.py": ROOT / "app.py",
            "ai.py": ROOT / "ai.py",
            "scanner.py": ROOT / "scanner.py",
            "config.json": CONFIG_PATH,
            "jobs.json": JOBS_PATH,
            "scan_status.json": STATUS_PATH,
            "requirements.txt": ROOT / "requirements.txt",
        }

        optional_files = {
            "favorites.json": FAVORITES_PATH,
            "history.json": HISTORY_PATH,
        }

        for file_name, file_path in required_files.items():
            if file_path.exists():
                file_size = file_path.stat().st_size

                st.success(
                    f"{file_name}: найден, {file_size} байт"
                )
            else:
                st.error(
                    f"{file_name}: отсутствует"
                )

        for file_name, file_path in optional_files.items():
            if file_path.exists():
                file_size = file_path.stat().st_size

                st.success(
                    f"{file_name}: найден, {file_size} байт"
                )
            else:
                st.info(
                    f"{file_name}: будет создан автоматически "
                    "при первом сохранении"
                )

    st.divider()

    diagnostic_section_3, diagnostic_section_4 = st.columns(
        2
    )

    with diagnostic_section_3:
        st.markdown("### 🌐 Источники заказов")

        configured_sources = cfg.get(
            "sources",
            [],
        )

        if not isinstance(
            configured_sources,
            list,
        ):
            configured_sources = []

        if not configured_sources:
            st.warning(
                "В config.json не найдены источники."
            )

        for source_index, source in enumerate(
            configured_sources
        ):
            if not isinstance(source, dict):
                continue

            source_name = str(
                source.get(
                    "name",
                    source.get(
                        "id",
                        f"Источник {source_index + 1}",
                    ),
                )
            )

            source_enabled = bool(
                source.get(
                    "enabled",
                    False,
                )
            )

            source_status_icon = (
                "✅"
                if source_enabled
                else "⏸️"
            )

            st.write(
                f"{source_status_icon} **{source_name}**"
            )

            source_type = source.get(
                "type",
                source.get(
                    "id",
                    "",
                ),
            )

            if source_type:
                st.caption(
                    f"Тип: {source_type}"
                )

            if source.get("url"):
                st.caption(
                    f"Адрес: {source.get('url')}"
                )

        if scan_status:
            with st.expander(
                "Последний отчёт сканера",
                expanded=False,
            ):
                st.json(
                    scan_status
                )

    with diagnostic_section_4:
        st.markdown("### 👤 Профиль")

        profile_display_name = str(
            profile.get(
                "display_name",
                "Не указано",
            )
        )

        profile_languages = profile.get(
            "languages",
            [],
        )

        profile_services = profile.get(
            "services",
            [],
        )

        st.write(
            f"**Имя:** {profile_display_name}"
        )

        if isinstance(
            profile_languages,
            list,
        ):
            st.write(
                "**Языки:** "
                + ", ".join(
                    str(language)
                    for language in profile_languages
                )
            )

        if isinstance(
            profile_services,
            list,
        ):
            st.write(
                "**Услуги:**"
            )

            for service in profile_services:
                st.write(
                    f"• {service}"
                )

        st.write(
            f"**Минимальная цена:** "
            f"{profile.get('minimum_price_eur', '—')} €"
        )

        st.write(
            f"**Предпочтительная цена:** "
            f"{profile.get('preferred_price_eur', '—')} €"
        )

        st.write(
            f"**Часовой пояс:** "
            f"{profile.get('timezone', '—')}"
        )

    st.divider()
    st.markdown("### 📦 Проверка библиотек")

    library_checks: dict[str, tuple[str, bool]] = {}

    try:
        import openai as openai_library

        library_checks["openai"] = (
            str(
                getattr(
                    openai_library,
                    "__version__",
                    "установлена",
                )
            ),
            True,
        )
    except Exception:
        library_checks["openai"] = (
            "не установлена",
            False,
        )

    try:
        import streamlit as streamlit_library

        library_checks["streamlit"] = (
            str(
                getattr(
                    streamlit_library,
                    "__version__",
                    "установлена",
                )
            ),
            True,
        )
    except Exception:
        library_checks["streamlit"] = (
            "не установлена",
            False,
        )

    try:
        import docx as docx_library

        library_checks["python-docx"] = (
            str(
                getattr(
                    docx_library,
                    "__version__",
                    "установлена",
                )
            ),
            True,
        )
    except Exception:
        library_checks["python-docx"] = (
            "не установлена",
            False,
        )

    try:
        import pypdf as pypdf_library

        library_checks["pypdf"] = (
            str(
                getattr(
                    pypdf_library,
                    "__version__",
                    "установлена",
                )
            ),
            True,
        )
    except Exception:
        library_checks["pypdf"] = (
            "не установлена",
            False,
        )

    library_columns = st.columns(
        len(library_checks)
    )

    for library_index, (
        library_name,
        library_information,
    ) in enumerate(
        library_checks.items()
    ):
        library_version, library_available = (
            library_information
        )

        library_columns[
            library_index
        ].metric(
            library_name,
            (
                library_version
                if library_available
                else "Отсутствует"
            ),
        )

    missing_libraries = [
        library_name
        for library_name, (
            _,
            library_available,
        ) in library_checks.items()
        if not library_available
    ]

    if missing_libraries:
        st.warning(
            "Не установлены библиотеки: "
            + ", ".join(missing_libraries)
        )
    else:
        st.success(
            "Все основные библиотеки установлены."
        )

    with st.expander(
        "Показать рекомендуемый requirements.txt",
        expanded=False,
    ):
        st.code(
            """streamlit>=1.36.0
openai>=1.0.0
python-docx>=1.1.0
pypdf>=4.0.0
feedparser>=6.0.11
requests>=2.31.0
""",
            language="text",
        )

    st.divider()
    st.markdown("### 🧹 Служебные действия")

    service_action_columns = st.columns(3)

    if service_action_columns[0].button(
        "Создать пустой favorites.json",
        use_container_width=True,
    ):
        if save_json(
            FAVORITES_PATH,
            get_favorites(),
        ):
            st.success(
                "favorites.json готов."
            )

    if service_action_columns[1].button(
        "Создать пустой history.json",
        use_container_width=True,
    ):
        if save_json(
            HISTORY_PATH,
            get_history(),
        ):
            st.success(
                "history.json готов."
            )

    if service_action_columns[2].button(
        "Перезагрузить приложение",
        use_container_width=True,
    ):
        st.rerun()


# ---------------------------------------------------------
# НИЖНЯЯ ЧАСТЬ СТРАНИЦЫ
# ---------------------------------------------------------

st.divider()

st.caption(
    "Фриланс-Агент помогает искать задания, создавать отклики, "
    "работать с текстами и проверять результат. Перед отправкой "
    "заказчику всегда самостоятельно проверяй готовую работу."
)
            
